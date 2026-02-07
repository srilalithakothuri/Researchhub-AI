from PyPDF2 import PdfReader
import os
from groq import Groq
from dotenv import load_dotenv

load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        raise Exception(f"PDF extraction error: {str(e)}")

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks"""
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start += chunk_size - overlap
    
    return chunks

def summarize_text(text: str, max_length: int = 500) -> str:
    """Generate summary of text using LLM"""
    try:
        # Truncate if too long
        text_to_summarize = text[:4000] if len(text) > 4000 else text
        
        messages = [
            {
                "role": "system",
                "content": "You are a research paper summarizer. Create concise, informative summaries."
            },
            {
                "role": "user",
                "content": f"Summarize this research paper in {max_length} words or less:\n\n{text_to_summarize}"
            }
        ]
        
        response = client.chat.completions.create(
            messages=messages,
            model="llama-3.3-70b-versatile",
            temperature=0.3,
            max_tokens=1024
        )
        
        return response.choices[0].message.content
    except Exception as e:
        return f"Summary generation failed: {str(e)}"

def extract_metadata_from_pdf(text: str) -> dict:
    """Extract title and authors from PDF text using LLM"""
    try:
        # Use first 2000 characters for metadata extraction
        text_sample = text[:2000]
        
        messages = [
            {
                "role": "system",
                "content": "Extract the title and authors from this research paper. Return in format: Title: [title]\nAuthors: [authors]"
            },
            {
                "role": "user",
                "content": text_sample
            }
        ]
        
        response = client.chat.completions.create(
            messages=messages,
            model="llama-3.3-70b-versatile",
            temperature=0.1,
            max_tokens=256
        )
        
        result = response.choices[0].message.content
        
        # Parse response
        title = "Unknown"
        authors = "Unknown"
        
        for line in result.split('\n'):
            if line.startswith("Title:"):
                title = line.replace("Title:", "").strip()
            elif line.startswith("Authors:"):
                authors = line.replace("Authors:", "").strip()
        
        return {"title": title, "authors": authors}
    except:
        return {"title": "Unknown", "authors": "Unknown"}
def generate_synthesized_report(summaries: list[str]) -> str:
    """Synthesize multiple research paper summaries into a coherent report"""
    try:
        combined_text = "\n\n---\n\n".join(summaries)
        
        messages = [
            {
                "role": "system",
                "content": "You are a senior research analyst. Synthesize the provided research summaries into a coherent, organized research report. Highlight common themes, conflicting findings, and unique contributions. Use professional language and Markdown formatting."
            },
            {
                "role": "user",
                "content": f"Please synthesize an executive research report from these individual paper summaries:\n\n{combined_text}"
            }
        ]
        
        response = client.chat.completions.create(
            messages=messages,
            model="llama-3.3-70b-versatile",
            temperature=0.4,
            max_tokens=2048
        )
        
        return response.choices[0].message.content
    except Exception as e:
        return f"Synthesis failed: {str(e)}"

def create_docx_report(content: str, output_path: str):
    """Save report content as a Word document"""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Research Synthesis Report', 0)
        
        # Simple markdown-ish to docx conversion
        for section in content.split('\n\n'):
            if section.startswith('# '):
                doc.add_heading(section.replace('# ', ''), level=1)
            elif section.startswith('## '):
                doc.add_heading(section.replace('## ', ''), level=2)
            elif section.startswith('### '):
                doc.add_heading(section.replace('### ', ''), level=3)
            else:
                doc.add_paragraph(section)
        
        doc.save(output_path)
    except Exception as e:
        print(f"Docx creation error: {str(e)}")
        raise e

def create_pdf_report(content: str, output_path: str):
    """Save report content as a PDF document"""
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('helvetica', '', 12)
        
        # Basic text wrapping
        for line in content.split('\n'):
            # Replace characters that might not be in standard Arial
            clean_line = line.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 10, clean_line)
            
        pdf.output(output_path)
    except Exception as e:
        print(f"PDF creation error: {str(e)}")
        raise e
